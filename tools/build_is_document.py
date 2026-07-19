from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/docx/01-IS-Documentation-completed.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=0):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))
    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:type"), "dxa")
    ind.set(qn("w:w"), str(indent_dxa))


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_width(table)
    for r, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(3)
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
            if r == 0:
                set_cell_shading(cell, "F1F3F4")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = str(text)
    style_table(table)
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)
    return table


def add_p(doc, text="", style=None, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    return p


def add_bullets(doc, items):
    for item in items:
        add_p(doc, item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        add_p(doc, item, style="List Number")


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after, color in (
        ("Heading 1", 20, 20, 6, "000000"),
        ("Heading 2", 16, 18, 6, "000000"),
        ("Heading 3", 14, 16, 4, "434343"),
    ):
        s = styles[name]
        s.font.name = "Arial"
        s.font.size = Pt(size)
        s.font.bold = False
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.line_spacing = 1.15

    for name in ("List Bullet", "List Number"):
        s = styles[name]
        s.font.name = "Arial"
        s.font.size = Pt(11)
        s.paragraph_format.left_indent = Inches(0.5)
        s.paragraph_format.first_line_indent = Inches(-0.25)
        s.paragraph_format.space_after = Pt(4)
        s.paragraph_format.line_spacing = 1.15


def title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Project Kalinga")
    r.font.name = "Arial"
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0, 0, 0)

    add_p(doc, "CEFMU Registry and Management System")
    add_p(doc, "Information System Documentation")
    add_p(doc, "Office: Social Technology Bureau")
    add_p(doc, "Division: Design Formulation Division")
    add_p(doc, "Date Started: May 2026")
    add_p(doc, "Target Completion: July 2026")
    add_p(doc, "Prepared for DSWD ICTMS review, QA readiness, VA readiness, and deployment documentation.")


def toc(doc):
    doc.add_page_break()
    add_p(doc, "Table of Contents", "Heading 1")
    items = [
        "1. Introduction",
        "2. Context Diagram and Architecture Narrative",
        "3. System Profile",
        "4. Business Process and Functional Coverage",
        "5. Access Control List",
        "6. Software and Hardware Requirements",
        "7. Data Dictionary",
        "8. Entity Relationship Description",
        "9. Security, Privacy, and Audit Controls",
        "10. Operations, Deployment, and Maintenance",
        "11. Manuals and Supporting Documents",
        "12. Revision History",
    ]
    add_bullets(doc, items)


def introduction(doc):
    add_p(doc, "1. Introduction", "Heading 1")
    add_p(
        doc,
        "Project Kalinga - CEFMU Registry and Management System is a web-based information system "
        "designed to help the Department of Social Welfare and Development record, monitor, validate, "
        "and report cases involving Child, Early, and Forced Marriage and Union. The system provides a "
        "single operational registry for authorized DSWD personnel and partner users so that case intake, "
        "case documentation, service delivery, monitoring, and reporting can be handled in a more timely, "
        "consistent, and accountable manner.",
    )
    add_p(
        doc,
        "The system responds to the operational need for structured documentation of CEFMU-related cases. "
        "Instead of relying on fragmented spreadsheets or manual consolidation, the application stores case "
        "records in a controlled data store, applies role-based access to case information, maintains a "
        "history of relevant user actions, and provides dashboard and report views for monitoring. It is "
        "intended to support program implementation while respecting confidentiality, data minimization, and "
        "the protection of personal and sensitive personal information.",
    )
    add_p(doc, "1.1 Purpose", "Heading 2")
    add_p(
        doc,
        "This documentation describes the purpose, scope, architecture, system profile, data structures, "
        "access controls, security controls, operating requirements, and supporting manuals for Project "
        "Kalinga. It is prepared as a reference for ICTMS review, technical assistance, quality assurance, "
        "vulnerability assessment readiness, deployment planning, and future maintenance.",
    )
    add_p(doc, "1.2 Scope", "Heading 2")
    add_p(
        doc,
        "The information system covers the core case management lifecycle from intake to monitoring, service "
        "documentation, and closure. It also covers administrative support functions such as user management, "
        "audit log review, report generation, public aggregate dashboard access, and limited offline operation "
        "for browser-side continuity.",
    )
    add_bullets(
        doc,
        [
            "Client intake and case registration for CEFMU-related cases.",
            "Case assessment, classification, and documentation of presenting problem, initial assessment, and plan of action.",
            "Family composition recording and case-related household information.",
            "Service delivery tracking, including service type, amount, date, and provider.",
            "Progress notes, follow-up notes, referral notes, closure notes, and related case narratives.",
            "Case status management, including active, closed, and reopened cases.",
            "Dashboard, report, and export features with audit logging for export activity.",
            "Role-based user management and authentication for DSWD personnel and authorized external users.",
            "Public dashboard access for aggregate, non-case-level statistics.",
        ],
    )
    add_p(doc, "1.3 Intended Users", "Heading 2")
    add_table(
        doc,
        ["User Type", "Primary Login Method", "Narrative Description"],
        [
            ["Admin", "Google OAuth or email/password", "DSWD system administrators who manage accounts, review logs, and oversee configuration and operational readiness."],
            ["Case Worker", "Google OAuth or email/password", "Authorized DSWD social workers who create cases, update case details, add services, and record progress notes."],
            ["FO User", "Email/password", "Field office or partner users who support case updating and service documentation within assigned coverage."],
            ["LGU Supervisor", "Email/password", "Local Government Unit supervisory users who monitor cases and updates within their jurisdiction."],
            ["CPU Monitor", "Email/password", "Monitoring users who require read-oriented dashboard and reporting access for oversight and coordination."],
            ["Public", "No login", "Unauthenticated users who may view only the public aggregate dashboard and FAQ, without case-level data."],
        ],
        [1.2, 1.6, 3.7],
    )


def architecture(doc):
    add_p(doc, "2. Context Diagram and Architecture Narrative", "Heading 1")
    add_p(
        doc,
        "Project Kalinga follows a serverless web application architecture. Users access a Vue 3 single-page "
        "application through a secure browser session. The frontend is built with Vite, deployed as static "
        "assets, and hosted on Vercel. The frontend communicates with a Google Apps Script web application "
        "over HTTPS. Google Apps Script performs authentication checks, authorization checks, validation, data "
        "read/write operations, caching, and audit logging. Google Sheets acts as the structured data store.",
    )
    add_p(
        doc,
        "The architecture separates the user interface from backend processing. The browser handles navigation, "
        "forms, offline queuing, cached reads, and user interaction. The Apps Script backend enforces the main "
        "security decisions before returning data or saving changes. This keeps privileged logic out of the "
        "browser and ensures that role and coverage checks occur on the server side before protected data is "
        "released.",
    )
    add_p(doc, "2.1 Logical Context Flow", "Heading 2")
    add_numbered(
        doc,
        [
            "An authorized user opens the Project Kalinga web application in a supported browser.",
            "The user authenticates through Google OAuth for DSWD personnel or through system-issued email/password credentials for external users.",
            "The frontend stores the session token locally and sends it with subsequent API requests.",
            "The Apps Script backend validates the token, identifies the user, checks role permissions, and applies coverage filters.",
            "Authorized requests read from or write to Google Sheets tables such as cases, users, services, progress_notes, family_members, sessions, and activity_log.",
            "Security-relevant actions are written to the audit log, and dashboard/report responses return only authorized or aggregate data.",
        ],
    )
    add_p(doc, "2.2 Component Summary", "Heading 2")
    add_table(
        doc,
        ["Component", "Technology", "Role in the System"],
        [
            ["Frontend", "Vue 3, Vite, Tailwind CSS, Pinia, Vue Router", "Provides case management, dashboards, reports, user management, FAQ, login, and offline-aware user interface."],
            ["Hosting", "Vercel static hosting and CDN", "Hosts compiled frontend assets over HTTPS and applies configured security headers."],
            ["Backend", "Google Apps Script web app", "Implements API actions, authentication, authorization, validation, caching, audit logging, and spreadsheet access."],
            ["Data Store", "Google Sheets", "Stores cases, users, services, notes, family members, sessions, lookups, locations, and activity logs."],
            ["Authentication", "Google OAuth and salted SHA-256 password authentication", "Supports domain-restricted DSWD login and admin-provisioned external user login."],
            ["Client Offline Support", "IndexedDB and service worker/PWA features", "Queues writes while offline and serves previously loaded data where browser cache is available."],
        ],
        [1.2, 1.6, 3.7],
    )


def profile(doc):
    add_p(doc, "3. System Profile", "Heading 1")
    add_table(
        doc,
        ["Profile Item", "Description"],
        [
            ["System Title", "Project Kalinga - CEFMU Registry and Management System"],
            ["System Type", "Web-based case registry, case management, monitoring, reporting, and administrative system."],
            ["Status", "For ICTMS technical assistance, QA readiness, VA readiness, and deployment documentation. Current implementation is prepared for review and hardening before production use."],
            ["Development Strategy", "Iterative development using Vue 3 for the frontend, Google Apps Script for backend services, Google Sheets for storage, Git/GitHub for version control, and Vercel for frontend deployment."],
            ["Computing Scheme", "Serverless single-page application architecture with static frontend hosting, Google Apps Script backend execution, and spreadsheet-backed data persistence."],
            ["Data Classification", "Contains personal and sensitive personal information and must be treated as confidential DSWD internal data."],
        ],
        [1.6, 4.9],
    )
    add_p(doc, "3.1 Hosting and Runtime Environment", "Heading 2")
    add_p(
        doc,
        "The frontend is compiled through Vite into static assets and served through Vercel. The backend is "
        "deployed as a Google Apps Script web app. The deployment model reduces the need to provision and "
        "maintain a traditional server, but it also requires strict control of script permissions, deployment "
        "versions, environment variables, Google OAuth configuration, and spreadsheet sharing permissions.",
    )
    add_p(doc, "3.2 System Inputs", "Heading 2")
    add_bullets(
        doc,
        [
            "User credentials, Google OAuth tokens, and backend session tokens.",
            "Client profile data, demographic details, location, classification, CEFMU type, admission mode, referral details, and intake information.",
            "Case assessment narratives, presenting problem, initial assessment, plan of action, remarks, and closure/reopening actions.",
            "Family member details, service records, progress notes, location updates, and report/export purpose declarations.",
            "Administrative data such as user roles, coverage assignments, password reset values, active status, and lookup values.",
        ],
    )
    add_p(doc, "3.3 System Outputs", "Heading 2")
    add_bullets(
        doc,
        [
            "Role-filtered case lists and case detail records.",
            "Dashboard summaries such as active/closed case counts, classification distribution, age bands, regional/provincial distribution, and monthly trends.",
            "Service, progress note, family composition, and location history views connected to each case.",
            "Exportable reports or CSV summaries subject to purpose declaration and audit logging.",
            "Administrative audit views covering activity logs, failed logins, lockouts, and export activity.",
            "Public aggregate statistics that do not expose personally identifiable case-level information.",
        ],
    )


def process_and_acl(doc):
    add_p(doc, "4. Business Process and Functional Coverage", "Heading 1")
    add_p(
        doc,
        "The system follows the operational flow of case registration, assessment, service tracking, "
        "monitoring, reporting, and administration. The case worker or authorized user begins by recording "
        "the client and case details. The case record then becomes the primary reference for services, notes, "
        "family composition, monitoring updates, and case status changes.",
    )
    add_p(doc, "4.1 Case Management Process", "Heading 2")
    add_numbered(
        doc,
        [
            "The user logs in and the system determines the user's role and coverage.",
            "The case worker creates a new case after reviewing the privacy notice and completing required intake fields.",
            "The system generates a unique case identifier using a CEFMU prefix, year-month component, and UUID fragment.",
            "The backend saves the case record, family member details, timestamps, and case worker information.",
            "Authorized users update the case, add services, add progress notes, save location details, close or reopen the case, and review related history.",
            "Dashboard and report views summarize the current case data while applying role and coverage restrictions.",
            "Exports and other relevant actions are logged for accountability.",
        ],
    )
    add_p(doc, "4.2 Functional Modules", "Heading 2")
    add_table(
        doc,
        ["Module", "Narrative Coverage"],
        [
            ["Login and Session Management", "Handles Google OAuth login for DSWD domain users, email/password login for external users, session token creation, expiry, logout, and password change requirements."],
            ["Dashboard", "Displays authorized statistical summaries, service totals, classification breakdowns, geographic distribution, and trend information."],
            ["Case Registry", "Provides search, filtering, viewing, creation, updating, closure, and reopening of CEFMU case records."],
            ["Case Detail", "Presents the case profile, assessment details, services, notes, family members, location information, and status actions."],
            ["Reports and Exports", "Provides summary tables/charts and export actions with purpose declaration and audit logging."],
            ["User Management", "Allows administrators to create users, set roles, assign coverage, set/reset passwords, and activate/deactivate accounts."],
            ["Audit Logs", "Allows administrators to review activity, failed login/lockout information, and export records."],
            ["Public Dashboard and FAQ", "Provides aggregate public information and help content without exposing case-level or personal data."],
        ],
        [1.6, 4.9],
    )
    add_p(doc, "5. Access Control List", "Heading 1")
    add_p(
        doc,
        "Access is based on authenticated identity, assigned role, and coverage fields such as LGU code, "
        "region, and province. The frontend hides modules that are not relevant to the user's role, while "
        "the backend performs authoritative permission checks before processing protected actions.",
    )
    add_table(
        doc,
        ["Function", "Admin", "Case Worker", "FO User", "LGU Supervisor", "CPU Monitor", "Public"],
        [
            ["Login", "Yes", "Yes", "Yes", "Yes", "Yes", "No"],
            ["Dashboard", "Yes", "Yes", "Yes", "Yes", "Yes", "Aggregate only"],
            ["View Cases", "Yes", "Coverage-based", "Coverage-based", "Coverage-based", "Read/report scope", "No"],
            ["Create Case", "No", "Yes", "No", "No", "No", "No"],
            ["Edit Case", "Yes", "Yes", "Yes", "Yes", "No", "No"],
            ["Close/Reopen Case", "Yes", "Yes", "Conditional", "Conditional", "No", "No"],
            ["Add Service", "Yes", "Yes", "Yes", "Yes", "No", "No"],
            ["Add/Edit Progress Notes", "Yes", "Yes", "Add where allowed", "Add where allowed", "No", "No"],
            ["Reports/Exports", "Yes", "Yes", "Yes", "Yes", "Yes", "No"],
            ["User Management", "Yes", "No", "No", "No", "No", "No"],
            ["Audit Logs", "Yes", "No", "No", "No", "No", "No"],
        ],
    )


def requirements(doc):
    add_p(doc, "6. Software and Hardware Requirements", "Heading 1")
    add_p(
        doc,
        "Because the system uses managed hosting and serverless backend execution, server requirements are "
        "primarily platform configuration requirements rather than physical server specifications. Client "
        "requirements focus on browser compatibility, stable connectivity, and device capacity sufficient for "
        "modern web application use.",
    )
    add_p(doc, "6.1 Server-Side Software Requirements", "Heading 2")
    add_bullets(
        doc,
        [
            "Vercel project configured for the Vue/Vite frontend build and deployment.",
            "Google Apps Script project deployed as a web app and authorized to access the required Google Sheets data store.",
            "Google Cloud OAuth client configured for DSWD Google sign-in and the production frontend origin.",
            "Google Sheets workbook containing the required tables/sheets and protected through appropriate sharing permissions.",
            "Environment variables for Google OAuth client ID and Google Apps Script web app URL.",
        ],
    )
    add_p(doc, "6.2 Client Software Requirements", "Heading 2")
    add_bullets(
        doc,
        [
            "Google Chrome 90 or later, Microsoft Edge 90 or later, Mozilla Firefox 90 or later, or Safari 15 or later.",
            "JavaScript enabled, local storage enabled, and IndexedDB support for offline queuing.",
            "Access to accounts.google.com when using Google OAuth login.",
            "Ability to download CSV/report exports when authorized.",
        ],
    )
    add_p(doc, "6.3 Hardware and Network Requirements", "Heading 2")
    add_table(
        doc,
        ["Area", "Minimum Requirement", "Recommended Requirement"],
        [
            ["Client device", "Desktop, laptop, tablet, or smartphone capable of running a modern browser.", "Desktop/laptop or modern tablet for extended case encoding and reporting work."],
            ["Memory", "2 GB RAM for basic browser use.", "4 GB RAM or higher for smoother multitasking."],
            ["Network", "Intermittent connectivity may allow limited offline queueing.", "Stable internet connection of at least 1 Mbps for normal use and synchronization."],
            ["Server hardware", "Not applicable due to serverless/static hosting model.", "Platform monitoring and access control review for Vercel, Apps Script, and Google Sheets."],
        ],
        [1.4, 2.2, 2.8],
    )


def data_dictionary(doc):
    add_p(doc, "7. Data Dictionary", "Heading 1")
    add_p(
        doc,
        "The data store is implemented in Google Sheets. Each sheet functions like a database table where "
        "the first row contains column names and each following row represents a record. The backend treats "
        "sheet names and header names as structured data contracts. The following data dictionary summarizes "
        "the primary fields and their intended meaning.",
    )
    tables = [
        (
            "7.1 cases",
            [
                ["case_id", "Primary identifier for a case. Generated with CEFMU prefix and date/UUID component."],
                ["date_intake", "Date when the case was registered or entered into the system."],
                ["status", "Current case status such as Active or Closed."],
                ["client_last, client_first, client_mi, suffix", "Client name fields."],
                ["birthdate, sex, age, civil_status, religion", "Core demographic fields."],
                ["ip_category, education, phone, occupation, income, philhealth_no", "Additional client and socioeconomic fields."],
                ["present_* and prov_* location fields", "Present and provincial address details by street, region, province, city/municipality, and barangay."],
                ["classification, cefmu_type, admission_mode", "Case classification and intake/admission categorization."],
                ["referred_by, referral_date", "Referral source and date if the case entered through referral."],
                ["presenting_problem, initial_assessment, plan_of_action, remarks", "Narrative assessment and intervention planning fields."],
                ["case_worker_email", "Email of the worker associated with the case."],
                ["date_closed, created_at, updated_at", "Lifecycle timestamps."],
            ],
        ),
        (
            "7.2 users",
            [
                ["email", "Primary identifier for system users."],
                ["display_name", "Name displayed in the system."],
                ["role", "Assigned role used for module visibility and backend authorization."],
                ["lgu_code, region, province", "Coverage fields used for filtering and authorization."],
                ["active", "Indicates whether the account may log in."],
                ["password_hash, salt", "Salted password hash values for email/password users."],
                ["must_change_password", "Forces password update after temporary password assignment."],
                ["failed_attempts, locked_until", "Login protection fields for lockout handling."],
                ["created_at", "Account creation timestamp."],
            ],
        ),
        (
            "7.3 services",
            [
                ["service_id", "Primary service record identifier."],
                ["case_id", "Foreign key linking the service to the case."],
                ["service_type", "Type/category of assistance or service provided."],
                ["amount", "Monetary value when applicable."],
                ["date_provided", "Date when the service was provided."],
                ["provided_by", "Provider or office that delivered the service."],
            ],
        ),
        (
            "7.4 progress_notes",
            [
                ["note_id", "Primary note identifier."],
                ["case_id", "Foreign key linking the note to a case."],
                ["date_note", "Date of the note."],
                ["note_type", "Progress, follow-up, referral, closure, or other configured note type."],
                ["content", "Narrative note content."],
                ["action_taken", "Actions already performed."],
                ["next_steps", "Recommended or planned next steps."],
                ["created_by, created_at", "Author and timestamp information."],
            ],
        ),
        (
            "7.5 family_members",
            [
                ["member_id", "Primary identifier for a family member record."],
                ["case_id", "Foreign key linking the family member to the case."],
                ["client_last, client_first, city_muni, province, region", "Client/location snapshot for reference."],
                ["name, birthdate, age, sex, relationship", "Family member identity and relationship details."],
                ["education, occupation, income", "Family member socioeconomic details."],
                ["created_at, updated_at", "Record timestamps."],
            ],
        ),
        (
            "7.6 sessions",
            [
                ["token", "Session token generated by the backend."],
                ["email", "User email associated with the session."],
                ["expires_at", "Session expiry timestamp."],
            ],
        ),
        (
            "7.7 activity_log",
            [
                ["timestamp", "Date/time when the activity occurred."],
                ["user_email", "User associated with the event."],
                ["action", "Action code such as LOGIN_PASSWORD, CREATE_CASE, UPDATE_CASE, ADD_SERVICE, EXPORT, or BLOCKED_*."],
                ["reference", "Related case ID, email, export reference, or other action-specific reference."],
                ["locale", "Locale or contextual value captured by the logging function when available."],
            ],
        ),
        (
            "7.8 lookups",
            [
                ["lookup_type", "Category of lookup, such as cefmu_type, note_type, or admission_mode."],
                ["value", "System value used in forms and records."],
                ["label", "User-facing label."],
                ["sort_order", "Display ordering value."],
            ],
        ),
    ]
    for heading, rows in tables:
        add_p(doc, heading, "Heading 2")
        add_table(doc, ["Field", "Description"], rows, [1.8, 4.7])


def erd_security_ops(doc):
    add_p(doc, "8. Entity Relationship Description", "Heading 1")
    add_p(
        doc,
        "Although implemented in Google Sheets, the system follows relational data design principles. The "
        "case record is the central entity. Services, progress notes, family members, and location records "
        "belong to a case through the case_id field. Users are related to cases through case_worker_email and "
        "to sessions/activity logs through email fields. Lookups standardize selectable values across forms.",
    )
    add_table(
        doc,
        ["Relationship", "Cardinality", "Description"],
        [
            ["users to sessions", "1 to many", "One user may have multiple session records over time. Each active session points back to one user email."],
            ["users to cases", "1 to many", "A case worker may be associated with many cases through case_worker_email."],
            ["cases to services", "1 to many", "A case may have zero or more service records."],
            ["cases to progress_notes", "1 to many", "A case may have zero or more notes documenting progress, referrals, follow-up, or closure."],
            ["cases to family_members", "1 to many", "A case may include multiple family member records."],
            ["users to activity_log", "1 to many", "User actions are logged with the user's email where available."],
            ["lookups to cases/notes/services", "reference", "Lookup values standardize dropdown options and reporting categories."],
        ],
        [1.7, 1.1, 3.7],
    )
    add_p(doc, "8.1 Key Constraints", "Heading 2")
    add_bullets(
        doc,
        [
            "case_id should be unique in the cases sheet.",
            "email should be unique in the users sheet.",
            "session tokens should be unique, temporary, and invalidated or expired when no longer valid.",
            "Services, notes, family members, and locations must reference an existing case_id.",
            "Role and coverage checks must be applied before a user can view, update, or export protected records.",
            "Public views must never expose personally identifiable case-level information.",
        ],
    )
    add_p(doc, "9. Security, Privacy, and Audit Controls", "Heading 1")
    add_p(
        doc,
        "The system processes sensitive personal information. Security and privacy controls must therefore be "
        "implemented both in the application and in the surrounding administrative procedures. The current "
        "implementation includes several technical controls and identifies areas that must remain part of "
        "deployment review, such as account provisioning, spreadsheet access, OAuth configuration, and export governance.",
    )
    add_p(doc, "9.1 Authentication and Session Controls", "Heading 2")
    add_bullets(
        doc,
        [
            "Google OAuth login is available for authorized DSWD personnel and is restricted to @dswd.gov.ph accounts.",
            "External users authenticate through administrator-provisioned email/password credentials.",
            "Passwords are stored as salted SHA-256 hashes rather than plaintext.",
            "Password complexity requires at least eight characters, an uppercase character, and a number.",
            "After five failed password attempts, the account is locked for fifteen minutes.",
            "Session tokens are generated by the backend, stored in the sessions sheet, and expire after eight hours.",
            "Logout invalidates the backend session token where possible.",
        ],
    )
    add_p(doc, "9.2 Authorization and Data Protection", "Heading 2")
    add_bullets(
        doc,
        [
            "The backend checks the requested action against the authenticated user's role before processing protected actions.",
            "Coverage fields such as LGU code, region, and province support role-filtered access to records.",
            "The frontend applies module visibility and route guards, while the backend remains the authoritative control point.",
            "Input sanitization is applied to reduce cross-site scripting risk from user-supplied values.",
            "Public dashboard output must be aggregate-only and must not reveal client-level identifiers.",
            "Export actions require a stated purpose and are logged for accountability.",
        ],
    )
    add_p(doc, "9.3 Audit Trail", "Heading 2")
    add_p(
        doc,
        "The activity log records significant events such as logins, failed logins, account lockouts, blocked "
        "actions, case creation, case updates, service additions, progress notes, case closure/reopening, "
        "location saves, user management actions, and exports. Administrators can review logs through the "
        "Audit Logs module. Export activity should be retained and reviewed according to DSWD policy because "
        "exports may carry higher privacy risk than ordinary viewing.",
    )
    add_p(doc, "10. Operations, Deployment, and Maintenance", "Heading 1")
    add_p(
        doc,
        "Operational readiness depends on both the deployed application and the administrative configuration "
        "around it. The system should be deployed using controlled environment variables, reviewed Apps Script "
        "deployment versions, restricted spreadsheet sharing, and documented account provisioning procedures. "
        "Any production deployment should be preceded by QA testing, vulnerability assessment review, and data "
        "privacy validation.",
    )
    add_p(doc, "10.1 Deployment Overview", "Heading 2")
    add_numbered(
        doc,
        [
            "Configure Google Sheets with the required sheets and header rows.",
            "Deploy Google Apps Script as a web app with the required permissions and access configuration.",
            "Configure the frontend environment variables for Google OAuth client ID and Apps Script web app URL.",
            "Build the Vue/Vite frontend and deploy the compiled static assets to Vercel.",
            "Configure Vercel security headers, including CSP, X-Frame-Options, X-Content-Type-Options, Referrer-Policy, and Permissions-Policy.",
            "Create initial administrator and test accounts, then validate role-based access and end-to-end case workflows.",
        ],
    )
    add_p(doc, "10.2 Maintenance Activities", "Heading 2")
    add_bullets(
        doc,
        [
            "Review access rights and deactivate users who no longer need access.",
            "Review failed login, lockout, export, and blocked action logs.",
            "Back up or export the Google Sheets data store according to approved records management procedures.",
            "Review lookup values and update categories only through authorized change control.",
            "Test application changes in staging before production deployment.",
            "Re-run QA and VA checks after security-sensitive changes.",
        ],
    )


def manuals_revision(doc):
    add_p(doc, "11. Manuals and Supporting Documents", "Heading 1")
    add_p(
        doc,
        "This IS documentation is supported by companion documents that provide operational instructions, "
        "deployment details, quality assurance evidence, and security readiness references. These documents "
        "should be maintained with the same versioning discipline as the system itself.",
    )
    add_table(
        doc,
        ["Document", "Purpose"],
        [
            ["User Manual", "Guides end users through login, dashboard use, case management, services, notes, reports, user management, audit logs, public dashboard, offline mode, FAQ, and troubleshooting."],
            ["Installation and Deployment Guide", "Documents setup, environment variables, Apps Script deployment, Vercel deployment, and operational configuration."],
            ["SQA Test Plan", "Defines test scope, approach, roles, environment, and acceptance criteria."],
            ["SQA Test Cases", "Provides scenario-level validation for authentication, role access, case workflows, dashboard/reporting, exports, and offline behavior."],
            ["SQA Compliance Checklist and Test Report", "Summarizes readiness against QA criteria and test execution results."],
            ["VA Security Compliance Checklist", "Supports vulnerability assessment and security control review."],
            ["User Stories", "Records role-based system requirements and expected user outcomes."],
        ],
        [2.0, 4.5],
    )
    add_p(doc, "12. Revision History", "Heading 1")
    add_table(
        doc,
        ["Version", "Date", "Author/Office", "Summary"],
        [
            ["1.0", "July 2026", "DSWD/STB", "Initial draft of Project Kalinga IS documentation."],
            ["1.1", "July 2026", "DSWD/STB with ICTMS preparation support", "Expanded narrative, architecture, ACL, data dictionary, security/privacy, deployment, and manual references for Google Docs submission."],
        ],
        [0.8, 1.0, 1.8, 2.9],
    )


def main():
    doc = Document()
    configure_styles(doc)
    title_block(doc)
    toc(doc)
    introduction(doc)
    architecture(doc)
    profile(doc)
    process_and_acl(doc)
    requirements(doc)
    data_dictionary(doc)
    erd_security_ops(doc)
    manuals_revision(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
